from docx import Document
import re

# Load the documents
doc_noref_path = './LUQ VII draft renewal proposal ver. 1.docx'
doc_ref_path = './LUQ VI Complete references March 1.docx'

doc_noref = Document(doc_noref_path)
doc_ref = Document(doc_ref_path)

# Function to extract text and page numbers from a document
def extract_text_and_page_numbers(doc):
    text_with_pages = []
    for i, paragraph in enumerate(doc.paragraphs):
        # Assuming a simple structure where page breaks can be identified
        if 'PAGE BREAK' in paragraph.text or i == 0: # Placeholder for actual page break detection
            page_number = len(text_with_pages) + 1  # Increment page number
            text_with_pages.append("")  # Start a new page
        text_with_pages[-1] += paragraph.text + "\n"  # Add text to the current page
    return text_with_pages

# Extract references and their text from the references document
def extract_full_references(doc):
    full_references = []
    for paragraph in doc.paragraphs:
        full_references.append(paragraph.text)
    return full_references

# Extract text with identified page numbers
text_with_pages_noref = extract_text_and_page_numbers(doc_noref)
full_references_list = extract_full_references(doc_ref)

# Check the first few entries to verify
text_with_pages_noref[:2], full_references_list[:2]


import pandas as pd

# Function to find partial references in the document
def find_partial_references(text_with_pages):
    partial_references = []
    # Regex pattern to match references like (Author et al. Year)
    pattern = re.compile(r'\([A-Za-z]+ et al\. \d{4}\)')
    for page_number, text in enumerate(text_with_pages, start=1):
        matches = pattern.findall(text)
        for match in matches:
            partial_references.append((match, page_number))
    return partial_references

# Function to match partial references to full references
def match_references(partial_refs, full_refs):
    matched_refs = []
    unmatched_refs = []  # List to keep track of unmatched references
    for partial_ref, page_number in partial_refs:
        matched = False  # Flag to indicate if a match was found
        year = re.search(r'\d{4}', partial_ref).group()
        for full_ref in full_refs:
            if year in full_ref and re.search(re.escape(partial_ref.split(" et al.")[0]), full_ref):
                matched_refs.append((partial_ref, full_ref, page_number))
                matched = True
                break
        if not matched:
            unmatched_refs.append((partial_ref, page_number))  # Add to unmatched if no match found
    return matched_refs, unmatched_refs

def find_name_and_name_year_references(text_with_pages):
    name_and_name_references = []
    # Specific regex pattern for "Name & 2nd Name Year"
    pattern = re.compile(r'\b([A-Za-z]+ & [A-Za-z]+ \d{4})\b')
    for page_number, text in enumerate(text_with_pages, start=1):
        matches = pattern.findall(text)
        for match in matches:
            name_and_name_references.append((match, page_number))
    return name_and_name_references

# Find partial references in the no-reference document
partial_references = find_partial_references(text_with_pages_noref)

# Match partial references to full references
matched_references, unmatched_references = match_references(partial_references, full_references_list)
df_matched_refs = pd.DataFrame(matched_references, columns=['Partial Reference', 'Full Reference', 'Page Number'])
df_unmatched_refs = pd.DataFrame(unmatched_references, columns=['Partial Reference',  'Page Number'])

# Create a DataFrame for the matched references
df_matched_refs = pd.DataFrame(matched_references, columns=['Partial Reference', 'Full Reference', 'Page Number'])

output_csv_path_unmatched = './unmatched_references7.csv'
# df_unmatched_refs.to_csv(output_csv_path_unmatched, index=False, encoding='utf-8-sig')

# Save the DataFrame to a CSV file
output_csv_path = './matched_references7.csv'
# df_matched_refs.to_csv(output_csv_path, index=False)

output_csv_path, df_matched_refs.head()

# Updated function to find partial references, considering multiple publications listed consecutively
def find_partial_references_updated(text_with_pages):
    partial_references_updated = []
    # Updated regex pattern to match references like (Author et al. Year) and consecutive listings
    pattern = re.compile(r'\((?:[A-Za-z]+ et al\. \d{4}(?:, )?)+\)')
    for page_number, text in enumerate(text_with_pages, start=1):
        matches = pattern.findall(text)
        for match in matches:
            # Extract individual references from the matched string
            individual_refs = re.findall(r'([A-Za-z]+ et al\. \d{4})', match)
            for ind_ref in individual_refs:
                partial_references_updated.append((ind_ref, page_number))
    return partial_references_updated

# Function to match partial references (updated to handle multiple references) to full references
def match_references_updated(partial_refs, full_refs):
    matched_refs_updated = []
    unmatched_refs_updated = []  # List to keep track of unmatched references
    for partial_ref, page_number in partial_refs:
        matched = False  # Flag to indicate if a match was found
        year = re.search(r'\d{4}', partial_ref).group()
        # Search for a match in the full references
        for full_ref in full_refs:
            if year in full_ref and re.search(re.escape(partial_ref.split(" et al.")[0]), full_ref):
                matched_refs_updated.append((partial_ref, full_ref, page_number))
                matched = True
                break  # Break after the first match to avoid duplicate entries for the same partial ref
        if not matched:
            unmatched_refs_updated.append((partial_ref, page_number))  # Add to unmatched if no match found
    return matched_refs_updated, unmatched_refs_updated


def generate_correct_pattern(partial_ref):
    # Split the reference by spaces, assuming " & " as the connector and the last part as the year
    parts = partial_ref.split(' ')
    year = parts[-1]  # The year is the last part
    authors = parts[:-1]  # Exclude the year for the authors part

    # Remove "&" or "and" if present; this assumes "&" is used as a connector
    if '&' in authors:
        authors.remove('&')

    # Escape individual components for authors and year
    escaped_authors = [re.escape(author) for author in authors]
    escaped_year = re.escape(year)

    # Reassemble the pattern, allowing for flexibility in whitespace
    # Since "&" or "and" is ignored in the full reference, we don't need to account for it in the pattern
    pattern = r'\s+'.join(escaped_authors) + r'\s+' + escaped_year

    return pattern


def generate_simplified_pattern(partial_ref):
    # Extract the year and author last names from the partial reference
    year = re.search(r'\d{4}', partial_ref).group()
    authors = re.split(r' & ', partial_ref[:-5])  # Remove the year and split the authors

    # Escape and prepare each author's last name for flexible matching
    pattern_authors = [re.escape(author.split(' ')[-1]) for author in authors]  # Focus on last names

    # Construct a regex pattern that allows for any content between the authors' last names and the year
    pattern = r'.*?'.join(pattern_authors) + r'.*?' + re.escape(year)

    return pattern

def flexible_match_references(partial_refs, full_refs):
    matched_refs = []
    unmatched_refs = []  # List to keep track of unmatched references
    for partial_ref, page_number in partial_refs:
        matched = False  # Flag to indicate if a match was found
        pattern = generate_simplified_pattern(partial_ref)
        regex = re.compile(pattern, re.IGNORECASE)
        for full_ref in full_refs:
            if regex.search(full_ref):
                matched_refs.append((partial_ref, full_ref, page_number))
                matched = True
                break  # Found a match, no need to continue looking
        if not matched:
            unmatched_refs.append((partial_ref, page_number))  # Add to unmatched if no match found
    return matched_refs, unmatched_refs

partial_ref = "Comita & Engelbrecht 2009"
pattern = generate_simplified_pattern(partial_ref)
print(f"Generated pattern: {pattern}")

# Compile the regex pattern with the IGNORECASE flag to make the search case-insensitive
regex = re.compile(pattern, re.IGNORECASE)

# Example test string, assuming a format in the full reference
test_string = "Comita, L. S. and B. M. Engelbrecht. 2009. Seasonal and spatial variation in water availability drive habitat associations in a tropical forest. Ecology 90:2755-2765."
match = regex.search(test_string)
if match:
    print("Match found:", match.group())
else:
    print("No match")

# Find updated partial references in the no-reference document
partial_references_updated = find_partial_references_updated(text_with_pages_noref)

matched_references_updated, unmatched_references_updated = match_references_updated(partial_references_updated, full_references_list)

# Convert both matched and unmatched references to DataFrames
df_matched_refs_updated = pd.DataFrame(matched_references_updated, columns=['Partial Reference', 'Full Reference', 'Page Number'])
df_unmatched_refs_updated = pd.DataFrame(unmatched_references_updated, columns=['Partial Reference', 'Page Number'])

# Save the DataFrames to CSV files
output_csv_path_matched_updated = './matched_references_updated7.csv'
# df_matched_refs_updated.to_csv(output_csv_path_matched_updated, index=False, encoding='utf-8-sig')

output_csv_path_unmatched_updated = './unmatched_references_updated7.csv'
# df_unmatched_refs_updated.to_csv(output_csv_path_unmatched_updated, index=False, encoding='utf-8-sig')




# Extracting "Name & 2nd Name Year" references specifically
name_and_name_references = find_name_and_name_year_references(text_with_pages_noref)

print(name_and_name_references)

# Matching these specific references to the full references list
# matched_name_and_name_references = match_references_comprehensive(name_and_name_references, full_references_list)
matched_name_and_name_references = flexible_match_references(name_and_name_references, full_references_list)

print("DONE")

matched_name_and_name_references, unmatched_name_and_name_references = flexible_match_references(name_and_name_references, full_references_list)

# Convert matched and unmatched references to DataFrames
df_matched_name_and_name = pd.DataFrame(matched_name_and_name_references, columns=['Partial Reference', 'Full Reference', 'Page Number'])
df_unmatched_name_and_name = pd.DataFrame(unmatched_name_and_name_references, columns=['Partial Reference', 'Page Number'])

df_matched = pd.concat([df_matched_refs_updated, df_matched_name_and_name], ignore_index=True)

# Save the DataFrames to CSV files
output_csv_path_matched = './matched_name_and_name_references.csv'
# df_matched.to_csv(output_csv_path_matched, index=False, encoding='utf-8-sig')

output_csv_path_unmatched = './unmatched_name_and_name_references.csv'
# df_unmatched_name_and_name.to_csv(output_csv_path_unmatched, index=False, encoding='utf-8-sig')

df_unmatched = pd.concat([df_unmatched_refs_updated, df_unmatched_name_and_name], ignore_index=True)

output_csv_path_name_and_name = './unmatched_references7.csv'
df_unmatched.to_csv(output_csv_path_name_and_name, index=False, encoding='utf-8-sig')


# Create and save a DataFrame for these specific matches
df_matched_name_and_name = pd.DataFrame(matched_name_and_name_references, columns=['Partial Reference', 'Full Reference', 'Page Number'])

# df_matched = pd.concat(df_matched_refs_updated, df_matched_name_and_name, ignore_index=True)
df_matched = pd.concat([df_matched_refs_updated, df_matched_name_and_name], ignore_index=True)

print(df_matched)

output_csv_path_name_and_name = './matched_references7.csv'
df_matched.to_csv(output_csv_path_name_and_name, index=False, encoding='utf-8-sig')
